import docx

def extract_study_plan():
    doc = docx.Document('../Test.docx')
    
    print("=== SEARCHING FOR STUDY PLAN ===")
    
    # First, search for "Study Plan" in all text
    print("\n1. Searching for 'Study Plan' text:")
    for i, para in enumerate(doc.paragraphs):
        if "Study Plan" in para.text:
            print(f"Found 'Study Plan' at line {i}: {para.text}")
            # Show context
            for j in range(max(0, i-3), min(len(doc.paragraphs), i+10)):
                marker = ">>>" if j == i else "   "
                print(f"{marker}{j:4d}: {doc.paragraphs[j].text.strip()}")
            print("---")
    
    # Then examine all tables thoroughly
    print("\n2. Examining all tables for year/semester structure:")
    for table_idx, table in enumerate(doc.tables):
        print(f"\n=== TABLE {table_idx + 1} ===")
        has_year_semester = False
        
        for row_idx, row in enumerate(table.rows):
            if not row.cells:
                continue
                
            # Extract all cells in this row
            row_cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            
            if not row_cells:
                continue
                
            # Look for year/semester patterns in any cell
            first_cell = row_cells[0] if row_cells else ""
            
            if any(pattern in first_cell.lower() for pattern in ['year', 'semester', 'csx', 'itx', 'ge']):
                print(f"Row {row_idx + 1}: {' | '.join(row_cells[:5])}")  # Limit columns for readability
                
                # If this looks like a year/semester header, show the full row
                if 'year' in first_cell.lower() or 'semester' in first_cell.lower():
                    print(f"  FULL: {' | '.join(row_cells)}")
                    has_year_semester = True
        
        if has_year_semester:
            print(f"^^^ TABLE {table_idx + 1} CONTAINS YEAR/SEMESTER STRUCTURE ^^^")
    
    print("\n=== SEARCH COMPLETE ===")

if __name__ == "__main__":
    extract_study_plan()
