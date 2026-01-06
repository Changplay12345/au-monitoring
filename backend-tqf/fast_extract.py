"""
Fast extraction module - No AI, uses regex patterns
Extracted algorithm from studyplan.py
"""
import re
from typing import List, Tuple
from models import Course, ProgramInfo, ParseResponse
from io import BytesIO
import docx
import PyPDF2


def extract_text_from_pdf(file_content: bytes) -> str:
    """Extract text from PDF file"""
    pdf_reader = PyPDF2.PdfReader(BytesIO(file_content))
    text_content = []
    for page in pdf_reader.pages:
        text = page.extract_text()
        if text:
            text_content.append(text)
    return '\n'.join(text_content)


def extract_text_from_docx(file_content: bytes) -> str:
    """Extract text from Word document"""
    doc = docx.Document(BytesIO(file_content))
    text_content = []
    
    # Extract text from paragraphs
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            text_content.append(paragraph.text.strip())
    
    # Extract text from tables
    # Handle cells with multiple lines (courses stacked in same cell)
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            
            # Check if any cell has newlines (multiple items stacked)
            max_lines = max(len(c.split('\n')) for c in cells) if cells else 1
            
            if max_lines > 1:
                # Split each cell by newlines and combine corresponding lines
                split_cells = [c.split('\n') for c in cells]
                for line_idx in range(max_lines):
                    line_parts = []
                    for cell_lines in split_cells:
                        if line_idx < len(cell_lines):
                            line_parts.append(cell_lines[line_idx].strip())
                        else:
                            line_parts.append('')
                    combined = ' '.join(p for p in line_parts if p)
                    if combined:
                        text_content.append(combined)
            else:
                # Single line per cell - join normally
                row_text = ' '.join(c for c in cells if c)
                if row_text:
                    text_content.append(row_text)
    
    return '\n'.join(text_content)


def extract_prerequisites_from_docx(file_content: bytes) -> dict:
    """
    Extract prerequisite mappings from Word document paragraphs.
    Returns dict: {course_code: prerequisite_string}
    Pattern: Course line followed by "Prerequisite: ..." line
    """
    doc = docx.Document(BytesIO(file_content))
    prerequisites = {}
    
    paras = [p.text.strip() for p in doc.paragraphs]
    
    for i, para in enumerate(paras):
        # Check if this line is a prerequisite line
        if para.lower().startswith('prerequisite'):
            # Look at previous non-empty line to find the course
            for j in range(i - 1, max(0, i - 5), -1):
                prev_para = paras[j]
                if not prev_para:
                    continue
                # Try to extract course code from previous line
                # Pattern: "CSX 3002 Object-Oriented Concepts and Programming \t3 (3-0-6) credits"
                course_match = re.match(r'^([A-Z]{2,4}\s*\d{4})', prev_para)
                if course_match:
                    course_code = course_match.group(1).replace(' ', '')
                    # Extract prerequisite content after "Prerequisite:" or "Prerequisites:"
                    prereq_match = re.match(r'Prerequisites?:\s*(.+)', para, re.IGNORECASE)
                    if prereq_match:
                        prereq_text = prereq_match.group(1).strip()
                        prerequisites[course_code] = prereq_text
                    break
    
    return prerequisites


def extract_credits_number(credits_str: str) -> int:
    """Extract credits number from string like '3 (3-0-6)'"""
    match = re.search(r"^(\d+)", credits_str)
    if match:
        return int(match.group(1))
    return 3  # default


def extract_semester_courses(text: str, year: int, semester: int) -> List[Tuple[str, str, int]]:
    """
    Extract courses for a specific year/semester
    Returns list of (code, title, credits) tuples
    """
    header_pattern = rf"Year\s*{year}[\s,]*Semester\s*{semester}"
    
    if not text:
        return []
    
    header_match = re.search(header_pattern, text, re.IGNORECASE)
    if not header_match:
        return []
    
    lines = text[header_match.start():].split('\n')
    data_rows = []
    found_table = False
    
    
    # For multi-line format parsing
    pending_codes = []
    pending_titles = []
    pending_credits = []
    
    for line in lines:
        line_stripped = line.strip()
        
        
        # Stop at Total line
        if re.search(r"Total", line, re.IGNORECASE):
            break
        
        # Look for table header
        if "Course Code" in line and "Course Title" in line and "Credits" in line:
            found_table = True
            continue
        
        if not found_table:
            continue
        
        # Handle Major Elective first (before other patterns)
        if "Major Elective" in line:
            elective_match = re.search(
                r"(One|Two|Three|Four|Five|Six|Seven|Eight|Nine|Ten|\d+)\s+Major Elective", 
                line, re.IGNORECASE
            )
            if elective_match:
                word_to_num = {
                    'one': 1, 'two': 2, 'three': 3, 'four': 4, 'five': 5,
                    'six': 6, 'seven': 7, 'eight': 8, 'nine': 9, 'ten': 10
                }
                num_str = elective_match.group(1).lower()
                try:
                    num = int(num_str)
                except ValueError:
                    num = word_to_num.get(num_str, 1)
                for _ in range(num):
                    data_rows.append(("", "Major Elective Course", 3, ""))
            continue
        
        # Handle Free Elective
        if "Free Elective" in line:
            elective_match = re.search(
                r"(One|Two|Three|Four|Five|Six|Seven|Eight|Nine|Ten|\d+)\s+Free Elective", 
                line, re.IGNORECASE
            )
            if elective_match:
                word_to_num = {
                    'one': 1, 'two': 2, 'three': 3, 'four': 4, 'five': 5,
                    'six': 6, 'seven': 7, 'eight': 8, 'nine': 9, 'ten': 10
                }
                num_str = elective_match.group(1).lower()
                try:
                    num = int(num_str)
                except ValueError:
                    num = word_to_num.get(num_str, 1)
                for _ in range(num):
                    data_rows.append(("", "Free Elective Course", 3, ""))
            continue
        
        # Find ALL courses on this line: CODE TITLE CREDITS pattern
        # Pattern matches: "CSX 3001 Fundamentals of Computer Programming 3 (3-0-6)"
        # Also handle "or CODE TITLE CREDITS" pattern
        
        # Check if line starts with "or" (case insensitive)
        is_or_course = line_stripped.lower().startswith('or ')
        search_line = line_stripped[3:].strip() if is_or_course else line_stripped
        
        # Also handle "CODE or CODE" pattern within line
        has_or_in_line = ' or ' in line_stripped.lower()
        
        # Check if NEXT line starts with "or" - means current line is start of OR group
        current_line_idx = lines.index(line) if line in lines else -1
        next_line_starts_with_or = False
        if current_line_idx >= 0 and current_line_idx + 1 < len(lines):
            next_line = lines[current_line_idx + 1].strip().lower()
            next_line_starts_with_or = next_line.startswith('or ')
        
        all_courses = re.findall(r"([A-Z]{2,4}\s*\d{4})\s+([^0-9]+?)\s+(\d+\s*\([\d\-]+\))", search_line)
        if all_courses:
            for i, course_match in enumerate(all_courses):
                # Remove spaces from course code to match AI extraction format
                code = course_match[0].strip().replace(" ", "")
                title = course_match[1].strip()
                credits = extract_credits_number(course_match[2].strip())
                # Mark as "or" if:
                # 1. Line started with "or"
                # 2. Multiple courses with "or" between them (not first one)
                # 3. Next line starts with "or" (this is first in OR group)
                # 4. Has "or" in line (all courses in that line are OR alternatives)
                or_flag = "or" if (is_or_course or (has_or_in_line and i > 0) or next_line_starts_with_or or has_or_in_line) else ""
                data_rows.append((code, title, credits, or_flag))
            continue
    
    return data_rows


def extract_program_info(text: str) -> ProgramInfo:
    """Extract program info from document text"""
    # Try to find program code pattern
    # Pattern: "Code" followed by number like "25330741100188"
    code_match = re.search(r"Code\s+(\d{10,})", text, re.IGNORECASE)
    if not code_match:
        code_match = re.search(r"Program\s*Code[:\s]*([A-Z0-9\-]+)", text, re.IGNORECASE)
    program_code = code_match.group(1) if code_match else "UNKNOWN"
    
    # Try to find program title
    # Pattern: "Program" followed by title like "Bachelor of Science Program in Computer Science (International Program)"
    title_match = re.search(r"Program\s+(Bachelor[^\n]+(?:\([^)]+\))?)", text, re.IGNORECASE)
    if title_match:
        program_title = title_match.group(1).strip()
    else:
        # Fallback: Look for "Bachelor of Science Program in" pattern
        title_match = re.search(r"(Bachelor\s+of\s+\w+\s+Program\s+in[^\n]+(?:\([^)]+\))?)", text, re.IGNORECASE)
        if title_match:
            program_title = title_match.group(1).strip()
        else:
            program_title = "Study Plan"
    
    # Clean up program title - remove extra whitespace
    program_title = ' '.join(program_title.split())
    
    # Try to find total credits
    credits_match = re.search(r"Total\s*(?:Credits?|หน่วยกิต)[:\s]*(\d+)", text, re.IGNORECASE)
    if not credits_match:
        # Try pattern like "132 Credits" or just count from structure
        credits_match = re.search(r"(\d{2,3})\s*Credits?", text, re.IGNORECASE)
    total_credits = int(credits_match.group(1)) if credits_match else 132
    
    return ProgramInfo(
        program_code=program_code,
        program_title=program_title,
        total_credits=total_credits
    )


def fast_extract_study_plan(file_content: bytes, filename: str) -> ParseResponse:
    """
    Fast extraction without AI - uses regex patterns
    Returns ParseResponse in the same format as Gemini extraction
    """
    # Extract text based on file type
    if filename.lower().endswith('.docx'):
        text = extract_text_from_docx(file_content)
        # Also extract prerequisites from DOCX
        prereq_map = extract_prerequisites_from_docx(file_content)
    else:
        text = extract_text_from_pdf(file_content)
        prereq_map = {}  # PDF prerequisite extraction not implemented yet
    
    if not text.strip():
        raise ValueError("Could not extract text from the uploaded file")
    
    # Extract program info
    program_info = extract_program_info(text)
    
    # Extract courses for all year/semester pairs
    year_sem_pairs = [
        (1, 1), (1, 2), (2, 1), (2, 2),
        (3, 1), (3, 2), (4, 1), (4, 2)
    ]
    
    courses: List[Course] = []
    
    for year, semester in year_sem_pairs:
        semester_courses = extract_semester_courses(text, year, semester)
        
        for course_data in semester_courses:
            # Unpack tuple (code, title, credits, or_flag)
            code = course_data[0]
            title = course_data[1]
            credits = course_data[2]
            or_flag = course_data[3] if len(course_data) > 3 else ""
            
            # Determine course type
            if "Major Elective" in title:
                course_type = "major_elective"
            elif "Free Elective" in title:
                course_type = "free_elective"
            else:
                course_type = "course"
            
            # Get prerequisite for this course (normalize code without space)
            code_normalized = code.replace(' ', '') if code else ''
            prereq = prereq_map.get(code_normalized, '')
            
            course = Course(
                year=year,
                semester=semester,
                course_code=code,
                course_title=title,
                credits=credits,
                prerequisite=prereq,
                or_flag=or_flag
            )
            courses.append(course)
    
    # Filter prerequisites to only include courses that exist in the plan
    valid_codes = {c.course_code.replace(' ', '') for c in courses if c.course_code}
    for course in courses:
        if course.prerequisite:
            # Extract course codes from prerequisite string
            prereq_codes = re.findall(r'([A-Z]{2,4}\s*\d{4})', course.prerequisite)
            valid_prereqs = []
            for prereq_code in prereq_codes:
                prereq_normalized = prereq_code.replace(' ', '')
                if prereq_normalized in valid_codes:
                    # Use normalized code (no spaces) to match AI format
                    valid_prereqs.append(prereq_normalized)
            # Update prerequisite to only include valid courses
            if valid_prereqs:
                course.prerequisite = ', '.join(valid_prereqs)
            else:
                course.prerequisite = ''
    
    return ParseResponse(
        program_info=program_info,
        courses=courses,
        session_id=None,
        graph=None
    )
