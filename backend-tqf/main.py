import os
import tempfile
import uuid
import asyncio
from typing import Dict, Any
from datetime import datetime, timedelta
from fastapi import FastAPI, File, UploadFile, HTTPException, BackgroundTasks, Response
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
import docx
import PyPDF2
from io import BytesIO
from dotenv import load_dotenv

# Load environment variables from .env file
load_dotenv()

from models import ParseResponse, ErrorResponse
from gemini_client import GeminiClient
from csv_utils import generate_csv, validate_and_clean_courses, generate_study_plan_graph
from fast_extract import fast_extract_study_plan

app = FastAPI(title="Study Plan Extractor", version="1.0.0")

# Enable CORS for frontend
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# In-memory storage for parsed data (in production, use Redis or database)
parsed_data_storage: Dict[str, ParseResponse] = {}
csv_storage: Dict[str, str] = {}
session_timestamps: Dict[str, datetime] = {}

# Initialize Gemini client
try:
    gemini_client = GeminiClient()
except ValueError as e:
    print(f"Warning: {e}")
    gemini_client = None


async def cleanup_expired_sessions():
    """
    Background task to clean up expired sessions (older than 1 hour)
    """
    while True:
        try:
            current_time = datetime.now()
            expired_sessions = [
                session_id for session_id, timestamp in session_timestamps.items()
                if current_time - timestamp > timedelta(hours=1)
            ]
            
            for session_id in expired_sessions:
                if session_id in parsed_data_storage:
                    del parsed_data_storage[session_id]
                if session_id in csv_storage:
                    del csv_storage[session_id]
                if session_id in session_timestamps:
                    del session_timestamps[session_id]
                    
                print(f"Cleaned up expired session: {session_id}")
            
        except Exception as e:
            print(f"Error during session cleanup: {e}")
        
        # Run cleanup every 30 minutes
        await asyncio.sleep(1800)


@app.on_event("startup")
async def startup_event():
    """Start background cleanup task"""
    asyncio.create_task(cleanup_expired_sessions())


def extract_text_from_docx(file_content: bytes) -> str:
    """Extract text from DOCX file including both paragraphs and tables"""
    doc = docx.Document(BytesIO(file_content))
    text_parts = []
    
    # Extract paragraph text
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            text_parts.append(paragraph.text.strip())
    
    # Extract table content with special formatting for year/semester headers
    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                if cell.text.strip():
                    row_text.append(cell.text.strip())
            
            if row_text:
                row_content = " | ".join(row_text)
                
                # Check if this is a year/semester header row
                first_cell = row_text[0].lower() if row_text else ""
                if any(pattern in first_cell for pattern in ['year', 'semester']):
                    # Add special formatting for year/semester headers
                    text_parts.append(f"\n=== {row_content} ===\n")
                else:
                    # Regular course row
                    text_parts.append(row_content)
    
    return "\n".join(text_parts)


def extract_text_from_pdf(file_content: bytes) -> str:
    """Extract text from PDF file"""
    pdf_reader = PyPDF2.PdfReader(BytesIO(file_content))
    text = ""
    for page in pdf_reader.pages:
        text += page.extract_text() + "\n"
    return text


@app.get("/")
async def root():
    """Health check endpoint"""
    return {"status": "healthy", "message": "Study Plan Extractor API"}


@app.post("/parse", response_model=ParseResponse)
async def parse_document(file: UploadFile = File(...)):
    """
    Parse uploaded DOCX or PDF file and extract study plan data
    """
    if not gemini_client:
        raise HTTPException(
            status_code=500, 
            detail="Gemini client not initialized. Please check GEMINI_API_KEY environment variable."
        )
    
    # Validate file type
    if not file.filename.lower().endswith(('.docx', '.pdf')):
        raise HTTPException(
            status_code=400,
            detail="Only DOCX and PDF files are supported"
        )
    
    try:
        # Read file content
        file_content = await file.read()
        
        # Extract text based on file type
        if file.filename.lower().endswith('.docx'):
            document_text = extract_text_from_docx(file_content)
        else:
            document_text = extract_text_from_pdf(file_content)
        
        if not document_text.strip():
            raise HTTPException(
                status_code=400,
                detail="Could not extract text from the uploaded file"
            )
        
        # Extract structured data using Gemini
        print("DEBUG: Calling gemini_client.extract_study_plan...")
        parse_response = gemini_client.extract_study_plan(document_text)
        print("DEBUG: Gemini extraction completed")
        
        # Validate and clean courses
        print("DEBUG: Validating and cleaning courses...")
        parse_response.courses = validate_and_clean_courses(parse_response.courses)
        print("DEBUG: Course validation completed")
        
        # Generate study plan graph
        print("DEBUG: About to call generate_study_plan_graph...")
        parse_response.graph = generate_study_plan_graph(parse_response.courses)
        print("DEBUG: Graph generation completed")
        
        # Generate unique ID for this parsing session
        session_id = str(uuid.uuid4())
        
        # Store parsed data with timestamp
        parsed_data_storage[session_id] = parse_response
        session_timestamps[session_id] = datetime.now()
        
        # Generate and store CSV
        csv_content = generate_csv(parse_response.courses)
        csv_storage[session_id] = csv_content
        
        # Create response with session_id
        parse_response.session_id = session_id
        
        return parse_response
        
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Failed to parse document: {str(e)}"
        )


@app.post("/parse-fast", response_model=ParseResponse)
async def parse_document_fast(file: UploadFile = File(...)):
    """
    Fast parse uploaded DOCX or PDF file using regex patterns (no AI)
    Much faster but does not extract prerequisites
    """
    # Validate file type
    if not file.filename.lower().endswith(('.docx', '.pdf')):
        raise HTTPException(
            status_code=400,
            detail="Only DOCX and PDF files are supported"
        )
    
    try:
        # Read file content
        file_content = await file.read()
        
        # Fast extraction using regex patterns
        print("DEBUG: Starting fast extraction (no AI)...")
        parse_response = fast_extract_study_plan(file_content, file.filename)
        print(f"DEBUG: Fast extraction completed - found {len(parse_response.courses)} courses")
        
        # Validate and clean courses
        print("DEBUG: Validating and cleaning courses...")
        parse_response.courses = validate_and_clean_courses(parse_response.courses)
        print("DEBUG: Course validation completed")
        
        # Generate study plan graph
        print("DEBUG: Generating study plan graph...")
        parse_response.graph = generate_study_plan_graph(parse_response.courses)
        print("DEBUG: Graph generation completed")
        
        # Generate unique ID for this parsing session
        session_id = str(uuid.uuid4())
        
        # Store parsed data with timestamp
        parsed_data_storage[session_id] = parse_response
        session_timestamps[session_id] = datetime.now()
        
        # Generate and store CSV
        csv_content = generate_csv(parse_response.courses)
        csv_storage[session_id] = csv_content
        
        # Create response with session_id
        parse_response.session_id = session_id
        
        return parse_response
        
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Failed to parse document: {str(e)}"
        )


@app.get("/csv/{session_id}")
async def download_csv(session_id: str):
    """
    Download CSV file for a specific parsing session
    """
    if session_id not in csv_storage:
        raise HTTPException(status_code=404, detail="CSV not found")
    
    csv_content = csv_storage[session_id]
    
    return Response(
        content=csv_content,
        media_type='text/csv',
        headers={'Content-Disposition': 'attachment; filename=study-plan.csv'}
    )


@app.get("/graph/{session_id}")
async def get_study_plan_graph(session_id: str):
    """Get study plan graph data for a specific parsing session"""
    if session_id not in parsed_data_storage:
        raise HTTPException(status_code=404, detail="Session not found")
    
    parse_response = parsed_data_storage[session_id]
    
    if not parse_response.graph:
        raise HTTPException(status_code=404, detail="Graph data not available")
    
    return parse_response.graph


@app.get("/program-info/{session_id}")
async def get_program_info(session_id: str):
    """Get program info for a specific parsing session"""
    if session_id not in parsed_data_storage:
        raise HTTPException(status_code=404, detail="Session not found")
    
    parse_response = parsed_data_storage[session_id]
    
    return parse_response.program_info


@app.delete("/cleanup/{session_id}")
async def cleanup_session(session_id: str):
    """
    Clean up stored data for a session
    """
    if session_id in parsed_data_storage:
        del parsed_data_storage[session_id]
    if session_id in csv_storage:
        del csv_storage[session_id]
    
    return {"message": "Session cleaned up successfully"}


if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8001)
